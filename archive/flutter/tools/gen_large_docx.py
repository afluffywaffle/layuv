#!/usr/bin/env python3
"""
Generate a ~120-page prose DOCX for the Flutter e-ink large-doc memory gate.
Uses python-docx. Output: /tmp/leamh_large_doc_test.docx
"""

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import random

LOREM_SENTENCES = [
    "The ancient manuscripts lay scattered across the stone table, each page a testament to the forgotten art of illumination.",
    "Scholars had debated for centuries whether the codex predated the great library or had been salvaged from its ashes.",
    "She traced the faded ink with a careful finger, reading the marginalia left by monks who had long since turned to dust.",
    "The vellum crackled softly as he turned each page, a sound like distant thunder in the otherwise silent scriptorium.",
    "Words that had once carried the weight of law and prophecy now rested quietly beneath her gaze.",
    "Every annotation told a story within the story — a reader's doubt, a scribe's correction, a priest's warning.",
    "The iron gall ink had eaten through the parchment in places, leaving lacunae that scholars filled with conjecture.",
    "He had devoted thirty years to the reconstruction of this text, and still three folios remained beyond his grasp.",
    "The binding was Coptic, the script Carolingian, yet the illuminations suggested a provenance neither could explain.",
    "Rain fell steadily outside the archive as she photographed the final recto, her breath fogging in the cold room.",
    "Translation is always betrayal, she reminded herself, and yet the alternative was silence.",
    "The palimpsest held two texts in collision: the original gospel and the later liturgical formulary scraped over it.",
    "Digital imaging had revealed ghost letters beneath the surface, a whisper from a text thought entirely lost.",
    "He argued that the colophon was a later addition, but the ink chemistry told a different story.",
    "Languages die the way fires die — first the bright flame, then the ember, then the imperceptible cooling of ash.",
    "Every manuscript is an act of faith: faith that there will be a reader, that the reader will understand, that understanding matters.",
    "The watermark placed the paper firmly in the fifteenth century, but the handwriting was not from any known workshop.",
    "She had catalogued eleven hundred manuscripts across nine countries and still felt humbled before a new acquisition.",
    "The gloss was more revealing than the text itself — argument embedded in commentary, resistance hiding in the margin.",
    "To annotate is to refuse passivity, to insist that the text is not finished, that reading is a kind of writing.",
    "The folio had survived fire, flood, deliberate suppression, and three hundred years of institutional neglect.",
    "He set down his magnifying glass and admitted, quietly, that the attribution would have to remain open.",
    "There is a peculiar intimacy in reading another person's marginal notes — you are sharing their surprise, their frustration.",
    "The provenance chain broke in 1943, as so many provenance chains break, in the chaos of a retreating army.",
    "Modern conservation had stabilized the pages, but the damage from the earlier rebinding was irreversible.",
    "She compared the letterforms against the atlas of scripts, column by column, ruling out one monastery after another.",
    "The rubrics were later additions, clearly — the ink sat on top of the main text rather than sinking beside it.",
    "A single red thread marked a passage the original reader had found remarkable; she found it remarkable too.",
    "He photographed the verso with raking light and there, unmistakably, were the impressions of a different text.",
    "The digital edition would make the manuscript accessible to any scholar with an internet connection — this was not nothing.",
]

CHAPTER_TITLES = [
    "On the Origins of Written Memory",
    "The Scriptorium and Its Discontents",
    "Marginalia as Resistance",
    "The Palimpsest Problem",
    "Ink, Vellum, and Impermanence",
    "Annotation and the Construction of Meaning",
    "Lost Libraries and the Shape of Absence",
    "The Colophon Speaks",
    "Digital Light on Ancient Darkness",
    "Reading as a Physical Act",
    "The Ethics of Reconstruction",
    "Provenance and Its Lacunae",
    "Script, Hand, and Workshop",
    "The Binding as Evidence",
    "Languages That Are Dying",
    "The Gloss and the Glossed",
    "Forgery and Its Discontents",
    "Three Disputed Folios",
    "On the Patience Required",
    "Afterword: The Unfinished Text",
]

def make_paragraph_text(seed, length="medium"):
    rng = random.Random(seed)
    n = {"short": 4, "medium": 8, "long": 14}[length]
    sentences = [rng.choice(LOREM_SENTENCES) for _ in range(n)]
    # Vary slightly so consecutive paragraphs differ
    result = " ".join(sentences)
    # Occasionally add a parenthetical
    if seed % 7 == 0:
        result += " (The footnote to this observation would require a monograph of its own.)"
    return result


def build_docx(out_path: str):
    doc = Document()

    # Page setup: A5-ish margins to get ~120 pages with body text
    section = doc.sections[0]
    section.page_width = Inches(5.83)   # A5 width
    section.page_height = Inches(8.27)  # A5 height
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.9)

    # Default Normal style: 11pt Literata-like (Times as fallback)
    normal_style = doc.styles["Normal"]
    normal_style.font.size = Pt(11)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = Pt(16)

    heading1_style = doc.styles["Heading 1"]
    heading1_style.font.size = Pt(16)
    heading1_style.font.bold = True

    heading2_style = doc.styles["Heading 2"]
    heading2_style.font.size = Pt(13)
    heading2_style.font.bold = True

    # Title page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_para.add_run("The Annotated Archive")
    run.bold = True
    run.font.size = Pt(22)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.add_run("Studies in Manuscript Culture, Digital Editions, and the\nPersistence of the Written Word")

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    author.add_run("A Test Document for Léamh Large-Document Pagination")
    author.runs[0].italic = True

    doc.add_page_break()

    # Table of contents placeholder
    doc.add_heading("Contents", level=1)
    for i, title in enumerate(CHAPTER_TITLES):
        doc.add_paragraph(f"{i + 1}.  {title}", style="Normal")
    doc.add_page_break()

    # Chapters
    seed = 1
    for chap_idx, chapter_title in enumerate(CHAPTER_TITLES):
        doc.add_heading(f"Chapter {chap_idx + 1}: {chapter_title}", level=1)

        # 3–5 sections per chapter
        n_sections = 3 + (chap_idx % 3)
        for sec_idx in range(n_sections):
            sec_title = f"§{chap_idx + 1}.{sec_idx + 1}  " + [
                "Introduction",
                "Background and Context",
                "The Evidence Examined",
                "Counterarguments",
                "Synthesis",
                "Conclusion",
            ][sec_idx % 6]
            doc.add_heading(sec_title, level=2)

            # 4–8 body paragraphs per section
            n_paras = 4 + (seed % 5)
            for p_idx in range(n_paras):
                length = ["short", "medium", "long"][seed % 3]
                text = make_paragraph_text(seed, length)
                para = doc.add_paragraph(text, style="Normal")
                # Occasionally justify
                if seed % 4 != 0:
                    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                # Occasionally bold a phrase in the para
                if seed % 11 == 0 and len(para.runs) > 0:
                    para.runs[0].bold = True
                seed += 1

            # Occasional block quote
            if seed % 5 == 0:
                quote = doc.add_paragraph(style="Normal")
                quote.paragraph_format.left_indent = Inches(0.4)
                quote.paragraph_format.right_indent = Inches(0.4)
                quote.add_run(
                    make_paragraph_text(seed + 100, "short")
                ).italic = True
                seed += 1

        # Page break between chapters
        doc.add_page_break()

    # Appendix
    doc.add_heading("Appendix: Selected Manuscript Descriptions", level=1)
    for i in range(30):
        doc.add_heading(f"MS {1000 + i * 17}: Codex {chr(65 + i % 26)}", level=2)
        for _ in range(3):
            doc.add_paragraph(make_paragraph_text(seed, "medium"), style="Normal")
            seed += 1

    doc.save(out_path)
    print(f"Saved: {out_path}")


if __name__ == "__main__":
    build_docx("/tmp/leamh_large_doc_test.docx")
